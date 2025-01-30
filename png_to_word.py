from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from PIL import Image
import os

def add_images_to_word(input_folder_path, output_file):
    # Create a new Word document
    doc = Document()

    # Set page layout to portrait and remove margins
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    # Get a list of all PNG files in the folder
    png_files = [f for f in os.listdir(input_folder_path) if f.endswith('.png')]

    if not png_files:
        print("No PNG files found in the folder.")
        return

    for png_file in png_files:
        # Get the full path to the image
        image_path = os.path.join(input_folder_path, png_file)

        # Open the image to check dimensions (optional, can resize if needed)
        with Image.open(image_path) as img:
            width, height = img.size
            print(f"Adding image: {png_file} (Width: {width}, Height: {height})")

        # Add the image to the Word document with text wrapping
        run = doc.add_paragraph().add_run()
        run.add_picture(image_path, width=Inches(8.5), height=Inches(11))

        # Add a page break after each image (except the last one)
        #if png_file != png_files[-1]:
            #doc.add_page_break()

    # Save the document
    doc.save(output_file)
    print(f"Images have been added to {output_file}")

if __name__ == "__main__":
    # Define the folder containing PNG images and the output Word file
    folder_path = input("Enter the folder path containing PNG images: ").strip()
    output_file = input("Enter the name of the output Word file (e.g., output.docx): ").strip()

    if not os.path.exists(folder_path):
        print("The specified folder does not exist.")
    else:
        add_images_to_word(folder_path, output_file)